"""Create sample .docx templates for development and testing."""
from pathlib import Path
from docx import Document


def create_test_case_template(output_path: str):
    doc = Document()
    doc.add_heading("测试用例文档", 0)

    doc.add_paragraph("项目名称: {{ project_name }}")
    doc.add_paragraph("产品信息: {{ product_info }}")
    doc.add_paragraph("版本信息: {{ version_info }}")
    doc.add_paragraph("生成日期: {{ generated_date }}")
    doc.add_paragraph("")

    doc.add_heading("测试用例列表", level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "用例标题"
    hdr[1].text = "测试步骤"
    hdr[2].text = "预期结果"

    doc.add_paragraph("")
    doc.add_paragraph("{% for case in test_cases %}")
    doc.add_paragraph("用例 {{ loop.index }}: {{ case.title }}")
    doc.add_paragraph("步骤: {{ case.steps | join('; ') }}")
    doc.add_paragraph("预期: {{ case.expected }}")
    doc.add_paragraph("---")
    doc.add_paragraph("{% endfor %}")

    doc.add_paragraph("")
    doc.add_paragraph("共 {{ total_cases }} 个测试用例")
    doc.save(output_path)
    print(f"Created: {output_path}")


def create_test_result_template(output_path: str):
    doc = Document()
    doc.add_heading("测试执行结果报告", 0)

    doc.add_paragraph("项目名称: {{ project_name }}")
    doc.add_paragraph("产品信息: {{ product_info }}")
    doc.add_paragraph("版本信息: {{ version_info }}")
    doc.add_paragraph("生成日期: {{ generated_date }}")
    doc.add_paragraph("")

    doc.add_heading("执行结果摘要", level=1)
    doc.add_paragraph("总用例数: {{ total_count }}")
    doc.add_paragraph("通过: {{ passed_count }}")
    doc.add_paragraph("失败: {{ failed_count }}")
    doc.add_paragraph("")

    doc.add_heading("详细结果", level=1)
    doc.add_paragraph("{% for r in test_results %}")
    doc.add_paragraph("用例: {{ r.case_title }}")
    doc.add_paragraph("结果: {{ r.result }}")
    doc.add_paragraph("实际结果: {{ r.actual }}")
    doc.add_paragraph("备注: {{ r.note }}")
    doc.add_paragraph("---")
    doc.add_paragraph("{% endfor %}")

    doc.save(output_path)
    print(f"Created: {output_path}")


if __name__ == "__main__":
    template_dir = Path(__file__).parent
    template_dir.mkdir(parents=True, exist_ok=True)

    create_test_case_template(str(template_dir / "example" / "test_case_template.docx"))
    create_test_result_template(str(template_dir / "example" / "test_result_template.docx"))
