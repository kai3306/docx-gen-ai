from pathlib import Path


def parse_text(content: str) -> str:
    """Parse plain text content."""
    return content.strip()


def parse_docx(file_path: str) -> str:
    """Parse .docx file and extract text content."""
    from docx import Document
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            tables.append(" | ".join(cells))
    return "\n".join(paragraphs + tables)


def parse_xlsx(file_path: str) -> str:
    """Parse .xlsx file and extract text content."""
    import openpyxl
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                lines.append(row_text)
    return "\n".join(lines)


def parse_file(file_path: str, original_filename: str) -> str:
    """Parse uploaded file based on its extension."""
    ext = Path(original_filename).suffix.lower()

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return parse_text(f.read())
    elif ext == ".md":
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return parse_text(f.read())
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext in (".xlsx", ".xls"):
        return parse_xlsx(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")
